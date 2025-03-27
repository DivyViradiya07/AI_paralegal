import os
import groq # type: ignore
from dotenv import load_dotenv # type: ignore
from docx import Document # type: ignore
import json
from PyPDF2 import PdfReader


load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
client = groq.Groq(api_key=api_key)


# FUNCTION FOR GIVING API CALL TO GROQ
def generate_legal_document(prompt):
    """Handles the API call to generate legal documents."""
    response = client.chat.completions.create(
        model="qwen-2.5-coder-32b",
        messages=[{"role": "system", "content": prompt}]
    )
    return response.choices[0].message.content

# Function to extract summarized text from PDF
def extract_text_from_pdf(pdf_path, max_pages=10, max_chars=5000):
    """Extracts limited text from the PDF to fit within token limits."""
    reader = PdfReader(pdf_path)
    extracted_text = []
    
    for i, page in enumerate(reader.pages[:max_pages]):  # Limit the number of pages
        text = page.extract_text()
        if text:
            extracted_text.append(text)
    
    full_text = "\n".join(extracted_text)
    
    # Truncate text if it exceeds the character limit
    return full_text[:max_chars] 



# PROMPTS FOR GENERATION OF CORE LEGAL DOCS
def generate_writ_petition(case_number, year, petitioner, respondent, court_name, jurisdiction, legal_grounds, relief_sought, supporting_documents):
    
    context_text = extract_text_from_pdf("D:\AI_Paralegal\SampleDocuments\Writ_petition.pdf")    
    prompt = f"""
    Refer to the following summarized writ petition format:    
    
    {context_text}
    
    Based on this structure, generate a new writ petition with the following details:
    
    Case Number: {case_number} of {year}
    Petitioner: {petitioner}
    Respondent: {respondent}
    Court: {court_name}
    Jurisdiction: {jurisdiction}
    Legal Grounds: {legal_grounds}
    Relief Sought: {relief_sought}
    Supporting Documents: {', '.join(supporting_documents)}

    Structure the writ petition with these sections:
    - Title & Case Details
    - Introduction & Parties Involved
    - Jurisdiction Basis
    - Facts of the Case
    - Legal Grounds & Precedents
    - Relief Sought
    - List of Supporting Documents
    - Prayer (Final Request to the Court)

    The document should be in a professional and structured legal format, using numbered paragraphs where necessary.
    """

    return generate_legal_document(prompt)




def generate_affidavit(case_number, year, petitioner, respondent, affiant_name, order_date, property_name, statement_of_facts):
    
    
    prompt = f"""
    Generate a formal legal affidavit in a professional tone with the following details:
    
    Case Number: {case_number} of {year}
    Petitioner: {petitioner}
    Respondent: {respondent}
    Affiant: {affiant_name}
    Order Date: {order_date}
    Property Name: {property_name}
    Statement of Facts: {statement_of_facts}

    Structure the affidavit in proper legal format with these sections:
    - Title & Case Details
    - Introduction (Details of the Affiant)
    - Statement of Facts
    - Benefits Offered (if applicable, such as rent & corpus details)
    - Compliance with Regulations
    - Oath & Affirmation
    - Signature & Notarization

    The document should be structured with numbered clauses where necessary, ensuring clarity and legal precision.
    """

    return generate_legal_document(prompt)



def generate_patent_application(application_number, year, inventor_name, assignee, title, field_of_invention, background, summary, claims, drawings_description):
    prompt = f"""
    Generate a formal patent application in a professional and structured legal format with the following details:

    Application Number: {application_number} of {year}
    Inventor: {inventor_name}
    Assignee: {assignee}
    Title of the Invention: {title}
    Field of Invention: {field_of_invention}
    Background: {background}
    Summary of Invention: {summary}
    Claims: {claims}
    Description of Drawings: {drawings_description}

    Structure the patent application with these sections:
    - Title & Application Details
    - Field of the Invention
    - Background of the Invention
    - Summary of the Invention
    - Detailed Description of the Invention
    - Claims (clearly defining the scope of protection)
    - Brief Description of Drawings (if applicable)
    - Abstract

    Ensure that the claims are structured clearly and precisely, with numbered points outlining each aspect of the invention.
    """

    return generate_legal_document(prompt)




def generate_annexure(case_number, year, petitioner, respondent, annexure_title, annexure_number, description, supporting_documents):
    prompt = f"""
    Generate a formal legal annexure in a professional and structured format with the following details:
    
    Case Number: {case_number} of {year}
    Petitioner: {petitioner}
    Respondent: {respondent}
    Annexure Title: {annexure_title}
    Annexure Number: {annexure_number}
    Description: {description}
    Supporting Documents: {', '.join(supporting_documents)}

    Structure the annexure document as follows:
    - Title (Annexure Number and Case Details)
    - Introduction (Brief description of the annexure and its relevance)
    - Detailed Description (Explaining the contents of the annexure)
    - List of Supporting Documents
    - Affirmation (Declaring the authenticity of the annexure)
    - Date & Signature of the Affiant or Legal Representative

    Ensure the document is formatted in a structured legal manner, maintaining clarity and professionalism.
    """

    return generate_legal_document(prompt)




def generate_witness_statement(case_number, year, court_name, witness_name, witness_details, statement):
    prompt = f"""
    Generate a formal witness statement in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Court: {court_name}
    Witness Name: {witness_name}
    Witness Details: {witness_details}
    Statement: {statement}

    Structure the witness statement as follows:
    - Title & Case Details
    - Witness Introduction (Name, Age, Occupation, Relationship to the Case)
    - Statement of Facts (Clear and chronological details of what the witness knows)
    - Affirmation (Witness swears under oath that the statement is true)
    - Signature & Date

    Ensure the document is formal and legally admissible.
    """

    return generate_legal_document(prompt)




def generate_exhibit(case_number, year, exhibit_number, exhibit_title, description, attached_documents):
    prompt = f"""
    Generate a formal exhibit document in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Exhibit Number: {exhibit_number}
    Exhibit Title: {exhibit_title}
    Description: {description}
    Attached Documents: {', '.join(attached_documents)}

    Structure the exhibit document as follows:
    - Title (Exhibit Number and Case Details)
    - Description of Exhibit (Purpose and relevance)
    - List of Attached Documents
    - Certification of Authenticity
    - Date & Signature of the Submitting Party

    Ensure proper legal formatting for use in court.
    """

    return generate_legal_document(prompt)




def generate_forensic_report(case_number, year, forensic_expert, forensic_field, report_summary, findings, conclusion):
    prompt = f"""
    Generate a formal forensic report in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Forensic Expert: {forensic_expert}
    Field of Analysis: {forensic_field}
    Report Summary: {report_summary}
    Findings: {findings}
    Conclusion: {conclusion}

    Structure the forensic report as follows:
    - Title & Case Details
    - Expert Introduction (Credentials and area of expertise)
    - Summary of the Investigation
    - Findings (Technical details, data analysis, results)
    - Conclusion (Expert opinion based on evidence)
    - Signature & Certification by the Forensic Expert

    Ensure technical accuracy and legal validity.
    """

    return generate_legal_document(prompt)




def generate_expert_opinion(case_number, year, expert_name, field_of_expertise, opinion_summary, detailed_opinion, supporting_references):
    prompt = f"""
    Generate a formal expert opinion document in a structured legal format with the following details:

    Case Number: {case_number} of {year}
    Expert Name: {expert_name}
    Field of Expertise: {field_of_expertise}
    Opinion Summary: {opinion_summary}
    Detailed Opinion: {detailed_opinion}
    Supporting References: {', '.join(supporting_references)}

    Structure the expert opinion document as follows:
    - Title & Case Details
    - Expert Introduction (Qualifications & Expertise)
    - Summary of Opinion
    - Detailed Explanation with Analysis
    - Supporting References (Case Laws, Precedents, Technical Reports)
    - Conclusion
    - Signature & Date

    Ensure that the document follows professional legal standards.
    """

    return generate_legal_document(prompt)



from docx import Document # type: ignore
from docx.shared import Pt # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # type: ignore
from docx.oxml import OxmlElement # type: ignore

def save_to_docx(content, filename, case_number=None, year=None):
    """Save formatted legal document content to a .docx file with enhanced formatting."""
    doc = Document()
    
    # Add a header with case details
    if case_number and year:
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.text = f"Case No: {case_number} | Year: {year}"
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = content.split("\n")  # Split content into lines

    for i, line in enumerate(lines, start=1):
        line = line.strip()
        if not line:
            continue  # Skip empty lines
        
        # Bold section titles
        if line.endswith(":"):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        
        # Numbered paragraphs
        elif len(line) > 50:  # Assuming long lines are main content
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {line}")  # Adding numbering
            run.font.size = Pt(11)
        
        else:
            doc.add_paragraph(line)  # Regular paragraph

    # Add a footer with a disclaimer
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.text = "This is a system-generated legal document. Verify with a legal professional."
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    full_path = os.path.join("D:\AI_Paralegal\Downloads", filename)  
    doc.save(full_path)


    # Save the document
    doc.save(filename)
    print(f"Document saved as {filename}")